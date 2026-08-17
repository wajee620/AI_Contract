"""Document parsing with per-page text-layer detection.

Native-text PDF pages are extracted directly (PyMuPDF); pages with no usable
text layer (scanned) are rendered to an image and sent to the vision-LLM OCR.
DOCX and TXT/MD are supported for convenience.
"""
from __future__ import annotations
import logging
from dataclasses import dataclass
from pathlib import Path
from typing import Callable, Optional

from app.llm.gateway import vision_ocr

log = logging.getLogger(__name__)

# Below this many extractable characters, a PDF page is treated as scanned.
TEXT_LAYER_MIN_CHARS = 40

ProgressCb = Optional[Callable[[str], None]]


@dataclass
class PageText:
    page: int  # 1-based
    text: str
    is_scanned: bool = False


def parse_document(path: str | Path, progress: ProgressCb = None) -> list[PageText]:
    path = Path(path)
    ext = path.suffix.lower()
    if ext == ".pdf":
        return _parse_pdf(path, progress)
    if ext == ".docx":
        return _parse_docx(path)
    if ext in (".txt", ".md"):
        return [PageText(page=1, text=path.read_text(encoding="utf-8", errors="replace"))]
    raise ValueError(f"Unsupported file type: {ext} (use PDF, DOCX, TXT or MD)")


def _parse_pdf(path: Path, progress: ProgressCb = None) -> list[PageText]:
    import fitz  # PyMuPDF

    pages: list[PageText] = []
    doc = fitz.open(str(path))
    try:
        total = doc.page_count
        for i, page in enumerate(doc):
            text = page.get_text("text").strip()
            if len(text) >= TEXT_LAYER_MIN_CHARS:
                pages.append(PageText(page=i + 1, text=text))
                continue
            if progress:
                progress(f"OCR (vision LLM) on scanned page {i + 1}/{total}")
            pix = page.get_pixmap(dpi=200)
            try:
                ocr_text = vision_ocr(pix.tobytes("png"))
            except Exception:
                log.exception("Vision OCR failed on page %s", i + 1)
                ocr_text = text  # whatever little we had
            pages.append(PageText(page=i + 1, text=ocr_text.strip(), is_scanned=True))
    finally:
        doc.close()
    return pages


def _parse_docx(path: Path) -> list[PageText]:
    import docx

    d = docx.Document(str(path))
    parts: list[str] = []
    for para in d.paragraphs:
        parts.append(para.text)
    for table in d.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return [PageText(page=1, text="\n".join(parts))]
